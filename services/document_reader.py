"""
Tender Document Reader — downloads and extracts text from tender documents.

Simplified port from tendly-main's TenderDocumentReader. Supports downloading
from procurement portals (EE, LV, PL, LT, GB, FR) and extracting text from
PDF, DOCX, XLSX, HTML, and TXT files. Falls back to AI summaries from the
database when downloads fail.
"""

import io
import re
import os
import time
import zipfile
import requests
from typing import Dict, List, Optional
from pathlib import Path

from core.database import get_session, Tender, TenderDetail, TenderDocuments

# Skip non-document files
SKIP_EXTENSIONS = {'.xml', '.json', '.sig', '.p7s', '.p7m', '.asice', '.bdoc'}
SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.xlsx', '.csv', '.html', '.htm', '.txt', '.zip'}

# Estonian portal endpoints
EE_VERSION_URL = "https://riigihanked.riik.ee/rhr/api/public/v1/procurement/{}/latest-version"
EE_DOCUMENTS_URL = "https://riigihanked.riik.ee/rhr/api/public/v1/proc-vers/{}/documents/general-info"
EE_TEMP_URL = "https://riigihanked.riik.ee/rhr/api/public/v1/proc-vers/{}/documents/{}/temp-url"


class TenderDocumentReader:
    """Handles fetching and extracting text from tender documents."""

    def __init__(self):
        self._session = None
        self._authenticated = False

    def _get_session(self) -> requests.Session:
        if self._session is None:
            self._session = requests.Session()
            self._session.headers.update({
                "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36",
                "Accept": "*/*",
            })
        return self._session

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def get_all_documents_from_db(self, tender_id: int) -> List[Dict]:
        """Fetch all documents for a tender from the database."""
        session = get_session()
        try:
            docs = (
                session.query(TenderDocuments)
                .filter(TenderDocuments.tender_id == tender_id)
                .all()
            )
            result = []
            for d in docs:
                ext = Path(d.file_name or d.document_name or "").suffix.lower()
                if ext in SKIP_EXTENSIONS:
                    continue

                ai_summary = (
                    d.ai_summary_en or d.ai_summary or ""
                )

                result.append({
                    "name": d.document_name or d.file_name or "",
                    "file_name": d.file_name or d.document_name or "",
                    "type": d.document_type or "",
                    "size": d.file_size or 0,
                    "ai_summary": ai_summary,
                    "web_url": d.web_url or "",
                    "procurement_doc_old_id": d.procurement_doc_old_id,
                })
            return result
        finally:
            session.close()

    def load_document_contents(self, tender_id: int, documents: List[Dict],
                               max_docs: int = 6) -> List[Dict]:
        """Load actual text content for tender documents.

        Tries to download and extract text. Falls back to AI summary if
        download fails.
        """
        loaded = []
        country_code = self._get_country_code(tender_id)

        for doc in documents[:max_docs]:
            doc_name = doc.get("name") or doc.get("file_name", "")
            file_name = doc.get("file_name", "") or doc_name
            ai_summary = doc.get("ai_summary", "")

            content = None
            try:
                content = self._download_and_extract(
                    tender_id, doc, country_code
                )
            except Exception as e:
                print(f"[DOC_READER] Failed to download {doc_name}: {e}")

            if content and len(content.strip()) > 50:
                loaded.append({
                    "name": doc_name,
                    "file_name": file_name,
                    "content": content[:8000],
                    "ai_summary": ai_summary,
                })
            elif ai_summary:
                loaded.append({
                    "name": doc_name,
                    "file_name": file_name,
                    "content": "",
                    "ai_summary": ai_summary,
                })

        return loaded

    def read_document(self, tender_id: int, document_name: str) -> Dict:
        """Read a single document by name. Returns {success, content, document_name, error}."""
        docs = self.get_all_documents_from_db(tender_id)
        # Find best match
        target = None
        for d in docs:
            if d["name"] == document_name or d["file_name"] == document_name:
                target = d
                break
        if not target:
            # Try substring match
            for d in docs:
                if document_name.lower() in (d["name"] or "").lower():
                    target = d
                    break

        if not target:
            return {"success": False, "content": "", "document_name": document_name,
                    "error": "Document not found"}

        country_code = self._get_country_code(tender_id)
        try:
            content = self._download_and_extract(tender_id, target, country_code)
            if content and len(content.strip()) > 50:
                return {"success": True, "content": content[:8000],
                        "document_name": target["name"], "error": None}
        except Exception as e:
            print(f"[DOC_READER] Download failed: {e}")

        # Fall back to AI summary
        if target.get("ai_summary"):
            return {"success": True, "content": target["ai_summary"],
                    "document_name": target["name"], "error": None}

        return {"success": False, "content": "", "document_name": target["name"],
                "error": "Could not extract document content"}

    # ------------------------------------------------------------------
    # Internal: country detection
    # ------------------------------------------------------------------

    def _get_country_code(self, tender_id: int) -> str:
        session = get_session()
        try:
            tender = session.query(Tender).filter(
                Tender.procurement_id == tender_id
            ).first()
            return tender.country_code if tender else "EE"
        finally:
            session.close()

    # ------------------------------------------------------------------
    # Internal: download and extract
    # ------------------------------------------------------------------

    def _download_and_extract(self, tender_id: int, doc: Dict,
                              country_code: str) -> Optional[str]:
        """Download document bytes and extract text."""
        doc_bytes = None

        if country_code == "EE":
            doc_bytes = self._download_estonian(tender_id, doc)
        elif doc.get("web_url"):
            doc_bytes = self._download_generic(doc["web_url"])

        if not doc_bytes:
            return None

        file_name = doc.get("file_name") or doc.get("name") or "unknown"
        return self.extract_text_from_bytes(doc_bytes, file_name)

    def _download_estonian(self, tender_id: int, doc: Dict) -> Optional[bytes]:
        """Download from Estonian procurement portal API."""
        sess = self._get_session()
        try:
            # Authenticate if needed
            if not self._authenticated:
                sess.get("https://riigihanked.riik.ee/rhr/api/public/v1/procurement/1/latest-version",
                         timeout=10)
                self._authenticated = True

            # Get version ID
            resp = sess.get(EE_VERSION_URL.format(tender_id), timeout=15)
            if resp.status_code != 200:
                return None
            version_id = resp.json() if isinstance(resp.json(), int) else resp.json().get("id")
            if not version_id:
                return None

            # Get temp download URL
            doc_old_id = doc.get("procurement_doc_old_id")
            if not doc_old_id:
                return None

            resp = sess.get(EE_TEMP_URL.format(version_id, doc_old_id), timeout=15)
            if resp.status_code != 200:
                return None

            download_url = resp.text.strip().strip('"')
            if not download_url.startswith("http"):
                return None

            # Download
            resp = sess.get(download_url, timeout=30)
            if resp.status_code == 200 and len(resp.content) > 100:
                return resp.content
        except Exception as e:
            print(f"[DOC_READER] Estonian download error: {e}")
        return None

    def _download_generic(self, url: str) -> Optional[bytes]:
        """Download from a generic URL."""
        if not url or not url.startswith("http"):
            return None
        sess = self._get_session()
        try:
            resp = sess.get(url, timeout=30, allow_redirects=True)
            if resp.status_code == 200 and len(resp.content) > 100:
                ct = resp.headers.get("Content-Type", "")
                if "text/html" in ct and len(resp.content) < 5000:
                    return None  # Probably an error page
                return resp.content
        except Exception as e:
            print(f"[DOC_READER] Generic download error for {url}: {e}")
        return None

    # ------------------------------------------------------------------
    # Internal: text extraction
    # ------------------------------------------------------------------

    def extract_text_from_bytes(self, content: bytes, file_name: str) -> Optional[str]:
        """Extract text from document bytes based on file extension."""
        ext = Path(file_name).suffix.lower()

        # Detect format from magic bytes if extension is unknown
        if ext not in SUPPORTED_FORMATS:
            ext = self._detect_format(content)

        if ext == ".pdf":
            return self._extract_pdf_text(content)
        elif ext == ".docx":
            return self._extract_docx_text(content)
        elif ext == ".xlsx":
            return self._extract_xlsx_text(content)
        elif ext in (".html", ".htm"):
            return self._extract_html_text(content)
        elif ext in (".txt", ".csv"):
            try:
                return content.decode("utf-8", errors="replace")[:8000]
            except Exception:
                return None
        elif ext == ".zip":
            return self._extract_from_zip(content, file_name)
        return None

    def _detect_format(self, content: bytes) -> str:
        if content[:4] == b"%PDF":
            return ".pdf"
        if content[:4] == b"PK\x03\x04":
            return ".docx"  # Could be docx, xlsx, or zip
        if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return ".doc"
        return ".unknown"

    def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages[:10]:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n\n".join(text_parts)
                return text[:8000] if text else None
        except Exception:
            pass

        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages[:10]:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            text = "\n\n".join(text_parts)
            return text[:8000] if text else None
        except Exception:
            pass
        return None

    def _extract_docx_text(self, content: bytes) -> Optional[str]:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs[:50]:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract table content
            for table in doc.tables[:5]:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        text_parts.append(" | ".join(cells))
            text = "\n".join(text_parts)
            return text[:8000] if text else None
        except Exception:
            return None

    def _extract_xlsx_text(self, content: bytes) -> Optional[str]:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            text_parts = []
            for ws in wb.worksheets[:3]:
                for row in ws.iter_rows(max_row=30, values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        text_parts.append(" | ".join(cells))
            text = "\n".join(text_parts)
            return text[:8000] if text else None
        except Exception:
            return None

    def _extract_html_text(self, content: bytes) -> Optional[str]:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "header", "footer"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            return text[:8000] if text else None
        except Exception:
            try:
                text = content.decode("utf-8", errors="replace")
                text = re.sub(r"<[^>]+>", " ", text)
                text = re.sub(r"\s+", " ", text).strip()
                return text[:8000] if text else None
            except Exception:
                return None

    def _extract_from_zip(self, content: bytes, original_name: str) -> Optional[str]:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for name in zf.namelist():
                    ext = Path(name).suffix.lower()
                    if ext in SUPPORTED_FORMATS and ext != ".zip":
                        inner = zf.read(name)
                        result = self.extract_text_from_bytes(inner, name)
                        if result:
                            return result
        except Exception:
            pass
        return None
