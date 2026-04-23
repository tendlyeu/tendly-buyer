"""
Generate DOCX from user stories markdown files using python-docx.
Plain formatting — just converts markdown structure to Word.

Usage:
    python docs/user_stories/generate_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).parent

FILES = {
    "en": {
        "md": ROOT / "user_stories_buyer_tools_en.md",
        "docx": ROOT / "Tendly_Buyer_Tools_User_Stories_EN.docx",
    },
    "ee": {
        "md": ROOT / "user_stories_buyer_tools_ee.md",
        "docx": ROOT / "Tendly_Buyer_Tools_User_Stories_EE.docx",
    },
}


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    table_rows = []
    in_table = False
    in_blockquote = False
    bq_lines = []

    def flush_blockquote():
        nonlocal in_blockquote, bq_lines
        if bq_lines:
            doc.add_paragraph(" ".join(bq_lines), style="Normal")
            bq_lines = []
        in_blockquote = False

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # First row is header
        cols = len(table_rows[0])
        tbl = doc.add_table(rows=len(table_rows), cols=cols, style="Table Grid")
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                tbl.rows[i].cells[j].text = cell
                for p in tbl.rows[i].cells[j].paragraphs:
                    p.style.font.size = Pt(10)
        table_rows = []
        in_table = False

    for line in lines:
        stripped = line.strip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            # Skip separator rows like |---|---|
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells):
                continue
            in_table = True
            if in_blockquote:
                flush_blockquote()
            table_rows.append(cells)
            continue
        elif in_table:
            flush_table()

        # Blockquote
        if stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            if content:
                bq_lines.append(content)
            in_blockquote = True
            continue
        elif in_blockquote:
            flush_blockquote()

        # Horizontal rule
        if stripped == "---" or stripped == "***":
            continue

        # Empty line
        if not stripped:
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        # List items
        elif stripped.startswith("- **"):
            # Bold-prefixed list item like "- **Q1.1** question text"
            clean = stripped[2:]  # remove "- "
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)  # strip bold markers
            doc.add_paragraph(clean, style="List Bullet")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        # Italic line (like the footer)
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            doc.add_paragraph(stripped.strip("*"), style="Normal")
        # Normal paragraph
        else:
            # Strip inline markdown
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            doc.add_paragraph(clean, style="Normal")

    # Flush remaining
    if in_table:
        flush_table()
    if in_blockquote:
        flush_blockquote()

    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"Generated: {docx_path} ({size_kb:.0f} KB)")


def main():
    for lang, config in FILES.items():
        if config["md"].exists():
            md_to_docx(config["md"], config["docx"])
        else:
            print(f"Skipping {lang}: {config['md']} not found")


if __name__ == "__main__":
    main()
